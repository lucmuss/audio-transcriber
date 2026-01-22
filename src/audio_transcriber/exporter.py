"""
Export module for converting transcriptions to various formats.
"""

import json
import logging
from pathlib import Path
from typing import Any, Dict, Optional

logger = logging.getLogger(__name__)


class TranscriptionExporter:
    """
    Export transcriptions to various formats (DOCX, Markdown, LaTeX, PDF).
    """

    def __init__(self):
        """Initialize exporter."""
        self.supported_formats = {"docx", "md", "markdown", "latex", "tex"}

    def export(
        self,
        transcription_file: Path,
        output_file: Path,
        export_format: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """
        Export transcription to specified format.

        Args:
            transcription_file: Path to input transcription file
            output_file: Path to output file
            export_format: Target format (docx, md, latex)
            metadata: Optional metadata (title, author, date, etc.)

        Returns:
            Dictionary with export results

        Raises:
            ValueError: If format is not supported
            FileNotFoundError: If transcription file doesn't exist
        """
        export_format = export_format.lower()

        if export_format not in self.supported_formats:
            raise ValueError(
                f"Unsupported format: {export_format}. "
                f"Supported: {', '.join(self.supported_formats)}"
            )

        if not transcription_file.exists():
            raise FileNotFoundError(f"Transcription file not found: {transcription_file}")

        # Read transcription content
        try:
            content = transcription_file.read_text(encoding="utf-8")
        except Exception as e:
            logger.error(f"Failed to read transcription file: {e}")
            return {
                "status": "error",
                "error": f"Failed to read file: {e}",
                "input_file": str(transcription_file),
            }

        # Select export method based on format
        if export_format == "docx":
            result = self._export_docx(content, output_file, metadata)
        elif export_format in ("md", "markdown"):
            result = self._export_markdown(content, output_file, metadata)
        elif export_format in ("latex", "tex"):
            result = self._export_latex(content, output_file, metadata)
        else:
            return {
                "status": "error",
                "error": f"Export method not implemented for {export_format}",
            }

        return result

    def _export_docx(
        self, content: str, output_file: Path, metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Export to Microsoft Word (.docx) format.

        Args:
            content: Transcription text
            output_file: Output file path
            metadata: Optional metadata

        Returns:
            Export result dictionary
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return {
                "status": "error",
                "error": "python-docx package not installed",
                "suggestion": "pip install python-docx",
            }

        try:
            doc = Document()

            # Add title if provided
            if metadata and metadata.get("title"):
                title = doc.add_heading(metadata["title"], 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add metadata section
            if metadata:
                if metadata.get("author"):
                    p = doc.add_paragraph()
                    p.add_run("Author: ").bold = True
                    p.add_run(metadata["author"])

                if metadata.get("date"):
                    p = doc.add_paragraph()
                    p.add_run("Date: ").bold = True
                    p.add_run(metadata["date"])

                if metadata.get("duration"):
                    p = doc.add_paragraph()
                    p.add_run("Duration: ").bold = True
                    p.add_run(metadata["duration"])

                if metadata.get("language"):
                    p = doc.add_paragraph()
                    p.add_run("Language: ").bold = True
                    p.add_run(metadata["language"])

                # Add separator
                doc.add_paragraph("_" * 50)

            # Add transcription heading
            doc.add_heading("Transcription", level=1)

            # Parse content (handle JSON formats)
            paragraphs = self._parse_content(content)

            # Add content paragraphs
            for para_text in paragraphs:
                if para_text.strip():
                    p = doc.add_paragraph(para_text)
                    # Set font
                    for run in p.runs:
                        run.font.name = "Calibri"
                        run.font.size = Pt(11)

            # Create output directory if needed
            output_file.parent.mkdir(parents=True, exist_ok=True)

            # Save document
            doc.save(output_file)

            logger.info(f"Exported to DOCX: {output_file.name}")

            return {
                "status": "success",
                "output_file": str(output_file),
                "format": "docx",
                "size_bytes": output_file.stat().st_size,
            }

        except Exception as e:
            logger.error(f"DOCX export failed: {e}")
            return {
                "status": "error",
                "error": f"DOCX export failed: {e}",
            }

    def _export_markdown(
        self, content: str, output_file: Path, metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Export to Markdown (.md) format.

        Args:
            content: Transcription text
            output_file: Output file path
            metadata: Optional metadata

        Returns:
            Export result dictionary
        """
        try:
            lines = []

            # Add title
            if metadata and metadata.get("title"):
                lines.append(f"# {metadata['title']}\n")

            # Add metadata table
            if metadata:
                lines.append("## Metadata\n")
                if metadata.get("author"):
                    lines.append(f"- **Author:** {metadata['author']}")
                if metadata.get("date"):
                    lines.append(f"- **Date:** {metadata['date']}")
                if metadata.get("duration"):
                    lines.append(f"- **Duration:** {metadata['duration']}")
                if metadata.get("language"):
                    lines.append(f"- **Language:** {metadata['language']}")
                lines.append("\n---\n")

            # Add transcription section
            lines.append("## Transcription\n")

            # Parse and add content
            paragraphs = self._parse_content(content)
            for para in paragraphs:
                if para.strip():
                    lines.append(para + "\n")

            # Join all lines
            markdown_content = "\n".join(lines)

            # Create output directory if needed
            output_file.parent.mkdir(parents=True, exist_ok=True)

            # Write file
            output_file.write_text(markdown_content, encoding="utf-8")

            logger.info(f"Exported to Markdown: {output_file.name}")

            return {
                "status": "success",
                "output_file": str(output_file),
                "format": "markdown",
                "size_bytes": output_file.stat().st_size,
            }

        except Exception as e:
            logger.error(f"Markdown export failed: {e}")
            return {
                "status": "error",
                "error": f"Markdown export failed: {e}",
            }

    def _export_latex(
        self, content: str, output_file: Path, metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Export to LaTeX (.tex) format.

        Args:
            content: Transcription text
            output_file: Output file path
            metadata: Optional metadata

        Returns:
            Export result dictionary
        """
        try:
            lines = []

            # LaTeX preamble
            lines.append(r"\documentclass[11pt,a4paper]{article}")
            lines.append(r"\usepackage[utf8]{inputenc}")
            lines.append(r"\usepackage[T1]{fontenc}")
            lines.append(r"\usepackage{lmodern}")
            lines.append(r"\usepackage[margin=2.5cm]{geometry}")
            lines.append(r"\usepackage{parskip}")
            lines.append(r"\usepackage{hyperref}")
            lines.append("")

            # Title and metadata
            if metadata:
                if metadata.get("title"):
                    title = self._escape_latex(metadata["title"])
                    lines.append(f"\\title{{{title}}}")

                if metadata.get("author"):
                    author = self._escape_latex(metadata["author"])
                    lines.append(f"\\author{{{author}}}")

                if metadata.get("date"):
                    date = self._escape_latex(metadata["date"])
                    lines.append(f"\\date{{{date}}}")
                else:
                    lines.append(r"\date{\today}")
            else:
                lines.append(r"\title{Transcription}")
                lines.append(r"\author{}")
                lines.append(r"\date{\today}")

            lines.append("")
            lines.append(r"\begin{document}")
            lines.append("")
            lines.append(r"\maketitle")
            lines.append("")

            # Add metadata section if available
            if metadata and any(
                metadata.get(k) for k in ("duration", "language", "model")
            ):
                lines.append(r"\section*{Metadata}")
                lines.append(r"\begin{itemize}")

                if metadata.get("duration"):
                    duration = self._escape_latex(metadata["duration"])
                    lines.append(f"  \\item \\textbf{{Duration:}} {duration}")

                if metadata.get("language"):
                    language = self._escape_latex(metadata["language"])
                    lines.append(f"  \\item \\textbf{{Language:}} {language}")

                if metadata.get("model"):
                    model = self._escape_latex(metadata["model"])
                    lines.append(f"  \\item \\textbf{{Model:}} {model}")

                lines.append(r"\end{itemize}")
                lines.append("")

            # Add transcription section
            lines.append(r"\section*{Transcription}")
            lines.append("")

            # Parse and add content
            paragraphs = self._parse_content(content)
            for para in paragraphs:
                if para.strip():
                    escaped_para = self._escape_latex(para)
                    lines.append(escaped_para)
                    lines.append("")  # Blank line between paragraphs

            lines.append(r"\end{document}")

            # Join all lines
            latex_content = "\n".join(lines)

            # Create output directory if needed
            output_file.parent.mkdir(parents=True, exist_ok=True)

            # Write file
            output_file.write_text(latex_content, encoding="utf-8")

            logger.info(f"Exported to LaTeX: {output_file.name}")

            return {
                "status": "success",
                "output_file": str(output_file),
                "format": "latex",
                "size_bytes": output_file.stat().st_size,
                "note": "Compile with: pdflatex filename.tex",
            }

        except Exception as e:
            logger.error(f"LaTeX export failed: {e}")
            return {
                "status": "error",
                "error": f"LaTeX export failed: {e}",
            }

    def _parse_content(self, content: str) -> list:
        """
        Parse transcription content, handling various formats.

        Args:
            content: Raw transcription content

        Returns:
            List of paragraphs
        """
        # Try to parse as JSON first (for verbose_json, json formats)
        try:
            data = json.loads(content)

            # Handle verbose_json format
            if isinstance(data, dict) and "text" in data:
                text = data["text"]
            # Handle segments format
            elif isinstance(data, dict) and "segments" in data:
                text = " ".join(seg.get("text", "") for seg in data["segments"])
            # Handle list of segments
            elif isinstance(data, list):
                text = " ".join(seg.get("text", "") for seg in data if isinstance(seg, dict))
            else:
                text = str(data)

            # Split into paragraphs
            paragraphs = text.split("\n")
            return [p.strip() for p in paragraphs if p.strip()]

        except (json.JSONDecodeError, ValueError):
            # Not JSON, treat as plain text
            paragraphs = content.split("\n")
            return [p.strip() for p in paragraphs if p.strip()]

    def _escape_latex(self, text: str) -> str:
        """
        Escape special LaTeX characters.

        Args:
            text: Text to escape

        Returns:
            Escaped text safe for LaTeX
        """
        # Special characters that need escaping in LaTeX
        replacements = {
            "\\": r"\textbackslash{}",
            "&": r"\&",
            "%": r"\%",
            "$": r"\$",
            "#": r"\#",
            "_": r"\_",
            "{": r"\{",
            "}": r"\}",
            "~": r"\textasciitilde{}",
            "^": r"\textasciicircum{}",
        }

        for old, new in replacements.items():
            text = text.replace(old, new)

        return text
